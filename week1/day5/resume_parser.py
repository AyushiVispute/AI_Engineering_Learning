import os
from pathlib import Path
from dotenv import load_dotenv
from groq import Groq
import time
from pydantic import BaseModel

load_dotenv()
my_api_key=os.getenv("GROQ_API_KEY")

if not my_api_key:
    raise ValueError("Api Error")

client=Groq(api_key=my_api_key)

model="llama-3.3-70b-versatile"

job_description="""
Full Stack Developer — Job Description

Role: Full Stack Developer
Skills: React.js, JavaScript, Node.js, Python, Flask, REST APIs, SQL/SQLite, HTML, CSS

Develop responsive and scalable full-stack web applications.
Build interactive frontend interfaces using React.js, JavaScript, HTML, and CSS.
Develop backend services and REST APIs using Node.js/Express.js and Python/Flask.
Integrate frontend applications with backend APIs and databases.
Work with SQL/SQLite and MongoDB for data management.
Implement authentication, CRUD operations, API integration, and responsive UI.
Debug, test, and optimize applications for performance and reliability.
Use Git/GitHub for version control and follow software development best practices.
Collaborate with team members to develop, maintain, and deploy web applications.

ATS Keywords: Full Stack Development, React.js, JavaScript, Node.js, Express.js, Python, Flask, REST API, HTML, CSS, SQL, SQLite, MongoDB, Git, GitHub, Backend Development, Frontend Development, API Integration, Database Management.
"""
class JobD(BaseModel):
    role:str
    required_skills:list[str]
    preferred_skils:list[str]
    minimum_experience: float | None = None
    education_reqirement:list[str]
    responsibilities:list[str]
    
jobd_schema=JobD.model_json_schema()

system_prompt=f"""
you are an expert HR assistance.

your job is to analyze job description and extract structured information from them.

return ONLY valid JSON matching this schema:
{jobd_schema}
IMPORTANT:
Do Not return the schema itself
Do Not return fields like "properties","title" or "type"
Fill the Schema with Actual information extracted from the job description.

if minimum experience is not mentioend,return an empty list.
Do not invent information"""

user_prompt=f""" Analyze the following job Description :
{job_description} """

message_system={
    "role":"system",
    "content":system_prompt
}
message_user={
    "role":"user",
    "content" : user_prompt
}
response_format={
    "type":"json_object"
}
messages=[message_system,message_user]

response=client.chat.completions.create(model=model,messages=messages,response_format=response_format)


answer=response.choices[0].message.content
raw_json=answer


import json
job_data=json.loads(raw_json)

job=JobD(**job_data)

print(job.minimum_experience)
print(job.education_reqirement)


class MatchResult(BaseModel):
    score:float
    details:dict

class Experience(BaseModel):
    company: str|None=None
    role:str|None=None 
    role:str|None=None 
    duration:str|None=None 
    description:str|None=None 
    skills_useed:list[str]=[]

class Resume(BaseModel):
    name:str|None=None
    email:str|None=None
    phone:str|None=None
    
    total_experience_years:float|None=None
    
    skills:list[str]=[]
    experiences:list[Experience]=[]
    education:list[str]=[]
    projects:list[str]=[]
    certificatoions:list[str]=[]
    
resume_schema=Resume.model_json_schema()

def  final_score(job,resume):
    match_schema=MatchResult.model_json_schema()
    prompt = f"""
    You are an HR recruiter.

    Compare the candidate's resume with the job description.

    JOB DESCRIPTION:
    {job.model_dump_json(indent=2)}

    CANDIDATE RESUME:
    {resume.model_dump_json(indent=2)}
    Return JSON matching this schema:

    {match_schema}

    Give me:

    1. Candidate name
    2. Matching skills
    3. Missing important skills
    4. Whether experience requirement is met
    5. Overall match percentage from 0 to 100
    6. A short final verdict

    Keep the response concise and easy to read.
    """
    message={
        "role": "user",
        "content" : prompt
    }
    messages=[message]
    response_format={
        "type": "json_object"
    }
    response = client.chat.completions.create(model=model, messages=messages, response_format=response_format)
    data = json.loads(response.choices[0].message.content)
    return MatchResult(**data)
def parse_resume(resume_text):
    system_prompt = f"""
    You are an expert resume parser.

    Extract information from the resume based on its meaning,
    not only based on exact section headings.

    Different resumes may use different headings.

    For example:
    - Experience
    - Professional Experience
    - Work History
    - Employment
    - Internships

    These may all contain relevant experience.

    Skills may also appear in the skills section, work experience,
    internships or projects.

    Return ONLY valid JSON matching this schema:

    {resume_schema}

    Important rules:

    1. Do not invent information.
    2. If a value is not available, return null.
    3. If a list has no information, return an empty list.
    4. Include internships inside experiences.
    5. Extract skills mentioned across the entire resume.
    """
    user_prompt = f"""
    Parse the following resume:

    {resume_text}
    """
    message_system={
        "role" : "system",
        "content" : system_prompt
    }
    message_user={
        "role" : "user",
        "content" : user_prompt
    }
    messages=[message_system, message_user]
    response_format={
        "type": "json_object"
    }
    response=client.chat.completions.create(model=model, messages=messages, response_format=response_format)
    raw_output = response.choices[0].message.content
    data = json.loads(raw_output)
    resume = Resume(**data)
    return resume


from pypdf import PdfReader
from docx import Document
def read_pdf(file_path):
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text

def read_docx(file_path):
    document = Document(file_path)
    text = ""
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text += paragraph.text + "\n"
    
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text += cell.text + "\n"
    return text


def read_resume(file_path):
    if file_path.suffix.lower() == ".pdf":
        return read_pdf(file_path)
    elif file_path.suffix.lower() == ".docx":
        return read_docx(file_path)
    else:
        return None



# lets do it now
resume_folder = Path("resumes")
all_results=[]
for file_path in resume_folder.iterdir():
    #C:\Users\Pratyush\padho_with_pratyush\week1\day5\resumes\abhay resume new - Abhay Singh.pdf
    if file_path.suffix.lower() not in [".pdf", ".docx"]:
        continue
    print("\nProcessing:", file_path.name)
    resume_text = read_resume(file_path)
    parsed_resume=parse_resume(resume_text) # llm call1
    time.sleep(5)
    result = final_score(job, parsed_resume) #llm caLL2
    #score and details
    #acount chtgpt
    # request bhejna shhur krega millions
    #chattgot server jam ho jayega
    time.sleep(5)
    print("Score:", result.score)
    all_results.append({
        "name": parsed_resume.name,
        "score": result.score,
        "details": result.details
    })
all_results.sort(
    key=lambda candidate: candidate["score"],
    reverse=True
)
top_2 = all_results[:2]
worst_2 = all_results[-2:]


print("TOP 2 CANDIDATES")
for candidate in top_2:

    print(
        candidate["name"],
        "-",
        candidate["score"],
        "%"
    )

    print(candidate["details"])

print("LOWEST 2 CANDIDATES")
for candidate in worst_2:

    print(
        candidate["name"],
        "-",
        candidate["score"],
        "%"
    )
    print(candidate["details"])

